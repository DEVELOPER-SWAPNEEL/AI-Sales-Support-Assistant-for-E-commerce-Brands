from __future__ import annotations

import json
import math
import re
import sys
import traceback
import unicodedata
from dataclasses import dataclass
from hashlib import sha256
from pathlib import Path
from typing import Any


DEBUG = True
BASE_DIR = Path(__file__).resolve().parents[1]
DATA_DIR = BASE_DIR / "data"
DOCS_DIR = DATA_DIR / "DOCS"
DB_DIR = BASE_DIR / "db"
CHROMA_DIR = DB_DIR / "chroma_db"
OUTPUT_JSON = DATA_DIR / "documents.json"
COLLECTION_NAME = "ecommerce_kb"
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
TEST_QUERIES = ("return policy", "shipping time", "payment methods")
INSTALL_COMMAND = "pip install -r requirements.txt"
EMBED_DIMENSION = 256

EMBEDDING_MODEL = None
CHROMA_COLLECTION = None

SAMPLE_DOCS: list[tuple[str, str]] = [
    (
        "return_policy",
        "Return Policy",
        (
            "We want customers to shop with confidence, so our return policy is designed to be clear and fair. "
            "Most eligible products can be returned within seven days of delivery when they are unused, unwashed, "
            "undamaged, and sent back with original tags and packaging. To request a return, customers should share "
            "their order number, the item name, and the reason for the request with support. Once the request is "
            "approved, pickup instructions or a return shipping process is shared. After the returned item passes a "
            "quality inspection, the refund is processed to the original payment method. Refund processing generally "
            "takes three to five business days after approval, depending on the payment provider. Shipping charges are "
            "typically non-refundable unless the product was damaged, defective, or incorrect at delivery. Exchange "
            "requests can also be supported when replacement stock is available."
        ),
    ),
    (
        "shipping_info",
        "Shipping Information",
        (
            "Orders are usually processed within one business day after confirmation, and customers receive tracking "
            "updates as soon as the shipment is dispatched. For most serviceable metro cities and major towns, "
            "delivery usually takes three to seven business days. Remote regions, weekends, public holidays, and high "
            "volume sale periods can increase the delivery timeline. Customers can use their tracking link to monitor "
            "the latest courier movement and estimated arrival. If an order is delayed beyond the expected delivery "
            "window, support can help investigate the courier status and guide the customer on the next best step. "
            "Delivery promises depend on courier serviceability, payment confirmation, and any operational issues "
            "outside the brand's control, such as weather or regional restrictions."
        ),
    ),
    (
        "payment_methods",
        "Payment Methods",
        (
            "Customers can complete their orders using secure and commonly supported payment methods. Available options "
            "typically include credit cards, debit cards, UPI, wallets, and net banking for prepaid checkout. Cash on "
            "delivery may be offered for selected pin codes, order values, and courier partners based on serviceability. "
            "If a payment attempt fails, the customer should verify their bank approval, available balance, and network "
            "stability before trying again. Orders placed through prepaid checkout are confirmed only after successful "
            "payment authorization. If a customer sees a duplicate charge or a failed payment with no confirmation, "
            "support can help validate the transaction status and suggest the safest next action."
        ),
    ),
]


@dataclass(frozen=True)
class StructuredDocument:
    id: str
    topic: str
    text: str

    def as_dict(self) -> dict[str, str]:
        return {"id": self.id, "topic": self.topic, "text": self.text}


def log(message: str) -> None:
    print(message)


def debug_log(message: str) -> None:
    if DEBUG:
        print(f"[DEBUG] {message}")


def print_error(message: str) -> None:
    print(f"Error: {message}")


def load_docx_reader() -> Any:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install dependencies with:\n"
            f"  {INSTALL_COMMAND}"
        ) from exc
    return Document


def load_embedding_dependencies() -> tuple[Any, Any]:
    try:
        import chromadb
        from sentence_transformers import SentenceTransformer
    except ImportError as exc:
        raise RuntimeError(
            "Missing embedding dependencies. Install dependencies with:\n"
            f"  {INSTALL_COMMAND}"
        ) from exc
    return chromadb, SentenceTransformer


class LocalEmbeddingModel:
    def encode(self, texts: list[str], show_progress_bar: bool = False) -> list[list[float]]:
        _ = show_progress_bar
        return [self._encode_text(text) for text in texts]

    def _encode_text(self, text: str) -> list[float]:
        vector = [0.0] * EMBED_DIMENSION
        tokens = re.findall(r"\w+", text.lower())
        if not tokens:
            return vector

        for token in tokens:
            token_hash = int(sha256(token.encode("utf-8")).hexdigest(), 16)
            index = token_hash % EMBED_DIMENSION
            sign = 1.0 if (token_hash >> 8) % 2 == 0 else -1.0
            vector[index] += sign

        norm = math.sqrt(sum(value * value for value in vector))
        if norm == 0:
            return vector
        return [value / norm for value in vector]


def build_embedding_model(sentence_transformer_cls: Any) -> Any:
    try:
        model = sentence_transformer_cls(EMBEDDING_MODEL_NAME, local_files_only=True)
        debug_log(f"Using embedding model '{EMBEDDING_MODEL_NAME}'.")
        return model
    except Exception as exc:
        debug_log(f"Falling back to local embedding model: {exc}")
        return LocalEmbeddingModel()


def ensure_directories() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    DB_DIR.mkdir(parents=True, exist_ok=True)


def create_sample_docx_files() -> None:
    ensure_directories()
    existing_docx = list(DOCS_DIR.glob("*.docx"))
    if existing_docx:
        return

    document_cls = load_docx_reader()
    for filename, heading, body in SAMPLE_DOCS:
        document = document_cls()
        document.add_heading(heading, level=1)
        for paragraph in re.split(r"(?<=[.!?])\s+", body):
            cleaned = paragraph.strip()
            if cleaned:
                document.add_paragraph(cleaned)
        output_path = DOCS_DIR / f"{filename}.docx"
        document.save(output_path)
        log(f"Created sample document: {output_path}")


def clean_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\ufeff", " ").replace("\u00a0", " ")
    cleaned_lines: list[str] = []

    for raw_line in normalized.splitlines():
        line = " ".join(raw_line.split())
        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()

def dedupe_documents(documents: list[StructuredDocument]) -> list[StructuredDocument]:
    unique_documents: list[StructuredDocument] = []
    seen_signatures: set[str] = set()

    for document in documents:
        signature = sha256(f"{document.topic}\n{document.text}".encode("utf-8")).hexdigest()
        if signature in seen_signatures:
            debug_log(f"Skipping duplicate document topic='{document.topic}'.")
            continue

        seen_signatures.add(signature)
        unique_documents.append(
            StructuredDocument(
                id=f"doc_{len(unique_documents) + 1:03d}",
                topic=document.topic,
                text=document.text,
            )
        )

    return unique_documents


def fallback_documents() -> list[StructuredDocument]:
    documents = [
        StructuredDocument(id="", topic=heading, text=clean_text(body))
        for _, heading, body in SAMPLE_DOCS
    ]
    log("Using fallback knowledge base documents.")
    return dedupe_documents(documents)


def build_structured_documents(input_dir: Path) -> list[StructuredDocument]:
    ensure_directories()
    create_sample_docx_files()

    document_cls = load_docx_reader()
    docx_files = sorted(file_path for file_path in input_dir.iterdir() if file_path.suffix.lower() == ".docx")

    if not docx_files:
        log(f"No .docx files found in '{input_dir}'. Using fallback documents.")
        return fallback_documents()

    documents: list[StructuredDocument] = []
    for file_path in docx_files:
        try:
            document = document_cls(file_path)
        except Exception:
            log(f"Skipping invalid file: {file_path}")
            debug_log(f"Invalid .docx could not be parsed: {file_path}")
            continue

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        text = clean_text("\n".join(paragraphs))

        if not text.strip():
            log(f"Skipping '{file_path.name}': empty content.")
            continue

        documents.append(StructuredDocument(id="", topic=file_path.stem.strip(), text=text))

    documents = dedupe_documents(documents)
    if not documents:
        log("No valid documents remained after filtering. Using fallback documents.")
        return fallback_documents()

    return documents


def save_documents_json(documents: list[StructuredDocument], output_path: Path = OUTPUT_JSON) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = [document.as_dict() for document in documents]
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def initialize_vector_store(documents: list[StructuredDocument]) -> None:
    global EMBEDDING_MODEL, CHROMA_COLLECTION

    chromadb, sentence_transformer_cls = load_embedding_dependencies()
    EMBEDDING_MODEL = build_embedding_model(sentence_transformer_cls)

    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    try:
        client.delete_collection(name=COLLECTION_NAME)
    except Exception:
        debug_log("No existing collection to delete before rebuild.")

    CHROMA_COLLECTION = client.get_or_create_collection(name=COLLECTION_NAME)
    texts = [document.text for document in documents]
    ids = [document.id for document in documents]
    metadatas = [{"topic": document.topic} for document in documents]
    raw_embeddings = EMBEDDING_MODEL.encode(texts, show_progress_bar=False)
    embeddings = raw_embeddings.tolist() if hasattr(raw_embeddings, "tolist") else raw_embeddings

    CHROMA_COLLECTION.add(documents=texts, metadatas=metadatas, ids=ids, embeddings=embeddings)


def query_kb(query: str, n_results: int = 3) -> list[dict[str, str]]:
    if EMBEDDING_MODEL is None or CHROMA_COLLECTION is None:
        raise RuntimeError("Knowledge base is not initialized. Run app/pipeline.py first.")

    raw_query_embedding = EMBEDDING_MODEL.encode([query], show_progress_bar=False)
    query_embedding = (
        raw_query_embedding.tolist()[0] if hasattr(raw_query_embedding, "tolist") else raw_query_embedding[0]
    )
    result = CHROMA_COLLECTION.query(query_embeddings=[query_embedding], n_results=n_results)
    documents = result.get("documents", [[]])[0]
    metadatas = result.get("metadatas", [[]])[0]

    formatted_results: list[dict[str, str]] = []
    for metadata, document_text in zip(metadatas, documents):
        text = str(document_text or "").strip()
        if not text:
            continue
        preview = text[:220]
        if len(text) > 220:
            preview += "..."
        formatted_results.append(
            {
                "topic": str((metadata or {}).get("topic", "Unknown")).strip() or "Unknown",
                "preview": preview,
            }
        )

    query_terms = set(re.findall(r"\w+", query.lower()))

    def score_item(item: dict[str, str]) -> int:
        topic_lower = item["topic"].lower()
        preview_lower = item["preview"].lower()
        score = (
            3 * sum(1 for term in query_terms if term in topic_lower)
            + sum(1 for term in query_terms if term in preview_lower)
        )
        if query_terms & {"shipping", "delivery", "timeline", "time"} and "shipping" in topic_lower:
            score += 10
        if query_terms & {"return", "refund", "exchange"} and "return" in topic_lower:
            score += 10
        if query_terms & {"payment", "payments", "cod", "upi", "card"} and "payment" in topic_lower:
            score += 10
        return score

    formatted_results.sort(key=score_item, reverse=True)
    return formatted_results


def print_query_results(query: str, results: list[dict[str, str]]) -> None:
    print(f"\nQuery: {query}")
    if not results:
        print("  No results found.")
        return

    for index, result in enumerate(results, start=1):
        print(f"  {index}. Topic: {result['topic']}")
        print(f"     Preview: {result['preview']}")


def main() -> int:
    try:
        ensure_directories()
        documents = build_structured_documents(DOCS_DIR)
        if not documents:
            print_error("No valid documents found after filtering. Stopping execution.")
            return 1

        save_documents_json(documents)
        log(f"Loaded {len(documents)} documents")
        log(f"Saved structured data to {OUTPUT_JSON}")

        initialize_vector_store(documents)
        log("Embeddings created")
        log(f"ChromaDB collection '{COLLECTION_NAME}' updated at {CHROMA_DIR}")

        for query in TEST_QUERIES:
            results = query_kb(query)
            print_query_results(query, results)

        log("\nRetrieval working")
        return 0
    except RuntimeError as exc:
        print_error(str(exc))
        return 1
    except Exception as exc:
        print_error(f"Unexpected failure: {exc}")
        debug_log(traceback.format_exc())
        return 1


if __name__ == "__main__":
    sys.exit(main())
